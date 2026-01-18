"""Utilities for extracting readable text from document formats."""

import re
from collections.abc import Sequence
from pathlib import Path

from docx import Document
from docx.table import Table

PathLike = str | Path

__all__ = ["export_docx_to_markdown", "extract_docx_text"]


def extract_docx_text(docx_path: PathLike) -> str:
    """Return a markdown-friendly string extracted from a DOCX file."""
    path = Path(docx_path)
    if not path.exists():
        raise ValueError(f"DOCX file not found: {path}")

    document = Document(str(path))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        line = _paragraph_to_markdown(
            paragraph.text, paragraph.style.name if paragraph.style else ""
        )
        if line:
            lines.append(line)

    for table in document.tables:
        lines.extend(_table_to_markdown(table))

    text = "\n\n".join(lines).strip()
    if not text:
        raise ValueError(f"No readable text extracted from {path.name}")
    return text


def export_docx_to_markdown(docx_path: PathLike, output_path: PathLike) -> Path:
    """Write DOCX contents to a markdown file and return the destination path."""
    destination = Path(output_path)
    text = extract_docx_text(docx_path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text(text + "\n", encoding="utf-8")
    return destination


def _paragraph_to_markdown(text: str, style_name: str) -> str:
    clean_text = text.strip()
    if not clean_text:
        return ""

    style_name = (style_name or "").lower()
    if "heading" in style_name:
        level = _heading_level(style_name)
        return f"{'#' * level} {clean_text}"
    if "list" in style_name:
        return f"- {clean_text}"

    return clean_text


def _heading_level(style_name: str) -> int:
    match = re.search(r"(\d+)", style_name)
    if match:
        level = int(match.group(1))
        return max(1, min(level, 6))
    return 2


def _table_to_markdown(table: Table) -> Sequence[str]:
    rows: list[str] = []
    for row in table.rows:
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        cells = [cell for cell in cells if cell]
        if cells:
            rows.append(" | ".join(cells))
    return rows
